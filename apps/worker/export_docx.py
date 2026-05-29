import json
import sys
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_base_style(document):
    normal_style = document.styles['Normal']
    normal_style.font.name = 'SimSun'
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    normal_style.font.size = Pt(11)

    for style_name, font_size in [('Title', 22), ('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 12)]:
        style = document.styles[style_name]
        style.font.name = 'SimHei'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        style.font.size = Pt(font_size)


def configure_page(document, title):
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.text = title

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' PAGE '
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)


def add_cover_page(document, title, generated_at, export_metadata):
    project_name = str(export_metadata.get('projectName') or '【待补充项目名称】')
    bid_reference_no = str(export_metadata.get('bidReferenceNo') or '【待补充招标编号】')
    bidder_name = str(export_metadata.get('bidderName') or '【待补充投标单位】')

    for _ in range(5):
        document.add_paragraph('')

    cover_title = document.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_title.style = document.styles['Title']
    cover_title.add_run(title)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run('技术方案导出文档')
    subtitle.runs[0].font.size = Pt(14)

    for _ in range(6):
        document.add_paragraph('')

    meta_table = document.add_table(rows=4, cols=2)
    meta_table.style = 'Table Grid'
    meta_table.cell(0, 0).text = '文档名称'
    meta_table.cell(0, 1).text = title
    meta_table.cell(1, 0).text = '文档类型'
    meta_table.cell(1, 1).text = '技术方案'
    meta_table.cell(2, 0).text = '导出时间'
    meta_table.cell(2, 1).text = generated_at
    meta_table.cell(3, 0).text = '招标项目 / 投标单位'
    meta_table.cell(3, 1).text = f'{project_name} / {bidder_name}'

    extra_table = document.add_table(rows=2, cols=2)
    extra_table.style = 'Table Grid'
    extra_table.cell(0, 0).text = '招标编号'
    extra_table.cell(0, 1).text = bid_reference_no
    extra_table.cell(1, 0).text = '生成方式'
    extra_table.cell(1, 1).text = '易标 Web 工作台'


def add_outline_page(document, outline_sections):
    document.add_page_break()
    document.add_heading('目录', level=1)

    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = '打开 Word 后右键更新目录。'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)

    document.add_paragraph('')

    if not outline_sections:
        document.add_paragraph('当前没有可用目录。')
        return

    for index, section in enumerate(outline_sections, start=1):
        document.add_paragraph(f'{index}. {str(section.get("title") or "")}', style='List Number')
        for child_index, child in enumerate(section.get('children') or [], start=1):
            document.add_paragraph(f'{index}.{child_index} {str(child.get("title") or "")}', style='List Number 2')


def add_summary_tables(document, outline_sections, knowledge_references, chapters):
    document.add_page_break()
    document.add_heading('文档摘要', level=1)

    summary_table = document.add_table(rows=4, cols=2)
    summary_table.style = 'Table Grid'
    summary_table.cell(0, 0).text = '一级章节数'
    summary_table.cell(0, 1).text = str(len(outline_sections))
    summary_table.cell(1, 0).text = '正文章节数'
    summary_table.cell(1, 1).text = str(len(chapters))
    summary_table.cell(2, 0).text = '知识引用数'
    summary_table.cell(2, 1).text = str(len(knowledge_references))
    summary_table.cell(3, 0).text = '导出格式'
    summary_table.cell(3, 1).text = 'DOCX'

    if knowledge_references:
        document.add_paragraph('')
        document.add_heading('知识库引用摘要', level=2)
        ref_table = document.add_table(rows=1, cols=2)
        ref_table.style = 'Table Grid'
        ref_table.cell(0, 0).text = '序号'
        ref_table.cell(0, 1).text = '引用条目'
        for index, reference in enumerate(knowledge_references, start=1):
            row = ref_table.add_row().cells
            row[0].text = str(index)
            row[1].text = str(reference.get('title') or '')


def is_list_line(line):
    stripped = line.strip()
    return stripped.startswith('- ') or stripped.startswith('* ') or stripped.startswith('1. ') or stripped.startswith('2. ') or stripped.startswith('3. ')


def is_table_block(lines):
    if len(lines) < 2:
        return False
    if '|' not in lines[0] or '|' not in lines[1]:
        return False

    separator = lines[1].replace('|', '').replace(':', '').replace('-', '').replace(' ', '')
    return separator == '' and '-' in lines[1]


def parse_table_row(line):
    return [cell.strip() for cell in line.strip().strip('|').split('|')]


def add_markdown_table(document, lines):
    rows = [parse_table_row(line) for line in lines if '|' in line]
    if len(rows) < 2:
        return False

    header = rows[0]
    body_rows = rows[2:]
    column_count = len(header)
    table = document.add_table(rows=1, cols=column_count)
    table.style = 'Table Grid'

    for index, cell in enumerate(header):
        table.cell(0, index).text = cell

    for row_data in body_rows:
        row = table.add_row().cells
        for index in range(column_count):
            row[index].text = row_data[index] if index < len(row_data) else ''

    return True


def add_list_block(document, lines):
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if stripped[:2].isdigit() and stripped[1:3] == '. ':
            document.add_paragraph(stripped[3:].strip(), style='List Number')
            continue

        if stripped[:1].isdigit() and stripped[1:3] == '. ':
            document.add_paragraph(stripped[3:].strip(), style='List Number')
            continue

        text = stripped[2:].strip() if stripped.startswith(('- ', '* ')) else stripped
        document.add_paragraph(text, style='List Bullet')


def add_paragraphs(document, text):
    blocks = [block.strip() for block in str(text or '').split('\n\n') if block.strip()]
    if not blocks:
        document.add_paragraph('')
        return

    for block in blocks:
        lines = [line.rstrip() for line in block.split('\n') if line.strip()]
        if is_table_block(lines):
            if add_markdown_table(document, lines):
                continue

        if lines and all(is_list_line(line) for line in lines):
            add_list_block(document, lines)
            continue

        paragraph = document.add_paragraph(block)
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        paragraph.paragraph_format.line_spacing = 1.5


def infer_heading_level(chapter, index):
    chapter_id = str(chapter.get('id') or '')
    if chapter_id.count('-') >= 2:
        return min(chapter_id.count('-'), 3)
    return 1 if index == 1 or True else 1


def build_heading_number(index_map, heading_level):
    if heading_level == 1:
        index_map[0] += 1
        index_map[1] = 0
        index_map[2] = 0
        return f'{index_map[0]}'

    if heading_level == 2:
        if index_map[0] == 0:
            index_map[0] = 1
        index_map[1] += 1
        index_map[2] = 0
        return f'{index_map[0]}.{index_map[1]}'

    if index_map[0] == 0:
        index_map[0] = 1
    if index_map[1] == 0:
        index_map[1] = 1
    index_map[2] += 1
    return f'{index_map[0]}.{index_map[1]}.{index_map[2]}'


def add_chapter(document, chapter, index, index_map):
    heading_level = infer_heading_level(chapter, index)
    heading_number = build_heading_number(index_map, heading_level)
    document.add_heading(f'{heading_number} {str(chapter.get("title") or "未命名章节")}', level=heading_level)

    info_table = document.add_table(rows=0, cols=2)
    info_table.style = 'Table Grid'
    if chapter.get('summary'):
        row = info_table.add_row().cells
        row[0].text = '章节摘要'
        row[1].text = str(chapter.get('summary'))
    if chapter.get('generatedAt') or chapter.get('updatedAt'):
        row = info_table.add_row().cells
        row[0].text = '生成时间'
        row[1].text = str(chapter.get('generatedAt') or '未记录')
        row = info_table.add_row().cells
        row[0].text = '更新时间'
        row[1].text = str(chapter.get('updatedAt') or chapter.get('generatedAt') or '未记录')

    references = chapter.get('references') or []
    if references:
        row = info_table.add_row().cells
        row[0].text = '参考条目'
        row[1].text = '、'.join(str(reference.get('title') or '') for reference in references)

    if len(info_table.rows) == 0:
        info_table._element.getparent().remove(info_table._element)

    add_paragraphs(document, chapter.get('content') or '')


def add_section_break(document):
    document.add_section(WD_SECTION_START.NEW_PAGE)


def main():
    if len(sys.argv) < 2:
        raise SystemExit('missing output path')

    output_path = sys.argv[1]
    payload = json.loads(sys.stdin.read() or '{}')

    document = Document()
    title = str(payload.get('title') or '未命名技术方案')
    generated_at = str(payload.get('generatedAt') or datetime.now().isoformat())
    export_metadata = payload.get('exportMetadata') or {}
    outline_sections = payload.get('outlineSections') or []
    knowledge_references = payload.get('knowledgeReferences') or []
    chapters = payload.get('chapters') or []

    set_base_style(document)
    configure_page(document, title)
    add_cover_page(document, title, generated_at, export_metadata)
    add_outline_page(document, outline_sections)
    add_summary_tables(document, outline_sections, knowledge_references, chapters)
    add_section_break(document)

    document.add_heading('正文内容', level=1)
    index_map = [0, 0, 0]
    for index, chapter in enumerate(chapters, start=1):
        add_chapter(document, chapter, index, index_map)

    document.save(output_path)


if __name__ == '__main__':
    main()
