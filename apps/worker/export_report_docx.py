import json
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def set_base_style(document):
    normal_style = document.styles['Normal']
    normal_style.font.name = 'SimSun'
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    normal_style.font.size = Pt(11)


def add_table(document, table_data):
    headers = table_data.get('headers') or []
    rows = table_data.get('rows') or []
    if not headers:
        return

    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for index, header in enumerate(headers):
        table.cell(0, index).text = str(header)

    for row_data in rows:
        row = table.add_row().cells
        for index in range(len(headers)):
            row[index].text = str(row_data[index] if index < len(row_data) else '')


def add_items(document, items):
    for item in items:
        document.add_heading(str(item.get('title') or '未命名项'), level=2)
        for line in item.get('lines') or []:
            document.add_paragraph(str(line), style='List Bullet')


def main():
    if len(sys.argv) < 2:
        raise SystemExit('missing output path')

    output_path = sys.argv[1]
    payload = json.loads(sys.stdin.read() or '{}')

    document = Document()
    set_base_style(document)

    title = str(payload.get('title') or '报告')
    summary = str(payload.get('summary') or '')
    sections = payload.get('sections') or []

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)

    if summary:
        document.add_paragraph(summary)

    for section in sections:
        document.add_heading(str(section.get('heading') or '未命名章节'), level=1)
        if section.get('table'):
            add_table(document, section.get('table') or {})
        if section.get('items'):
            add_items(document, section.get('items') or [])

    document.save(output_path)


if __name__ == '__main__':
    main()
